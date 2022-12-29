#python3

"""
I didn't write this code, I just amend it slightly
requirements:
document==1.0
lxml==4.2.4
Pillow==5.2.0
python-docx==0.8.7
Reference: https://python-docx.readthedocs.io/en/latest/index.html
"""

from docx import Document
from docx.shared import Inches

import datetime
from pathlib import Path
import os

time = datetime.datetime.now()
output_time = time.strftime("%Y%m%d_%H%M_")


    
document = Document()

document.add_heading('Title: RBCs change volumes ', level=1) #Heading, level 1

p = document.add_paragraph('Peyman Obeidy, et al ')
p.add_run('Corresponding author: Email:').bold = True
p.add_run(' School of, Faculty of ')
p.add_run('et al.').italic = True

document.add_heading('Abstract ', level=2) #Heading, level 2
p = document.add_paragraph('The mechanical force defines the shape properties of of red blood cells. To pass through capillaries much narrower than the RBCs diameter and optimal regulation of dynamic deformability is required for RBC. ')

document.add_heading('Disclosures   ', level=2) #Heading, level 2

document.add_paragraph('The authors declare no conflict of interest. ', style='Intense Quote')

document.add_heading('Figure legends    ', level=2) #Heading, level 2

document.add_paragraph(
    'Figure 1. ', style='List Bullet')
document.add_paragraph(
    'Figure 2. ', style='List Bullet')
    
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save(str(output_time)+"test_abstract.docx")
print(40*'*')
print("Done!, a test word document has been generated, \and if you don't believe me check your root folder")